"""Report rendering: Markdown, HTML, PDF (WeasyPrint), DOCX (python-docx)."""
from __future__ import annotations

import io
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import bleach
import markdown as markdown_lib
from jinja2 import Environment, FileSystemLoader

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"

# Dataset column names/values flow unescaped into the markdown template, so the
# rendered HTML must be sanitized before the frontend renders it via
# dangerouslySetInnerHTML — a crafted CSV column like "<script>..." must not survive.
_ALLOWED_TAGS = [
    "p", "br", "hr", "h1", "h2", "h3", "h4", "ul", "ol", "li", "strong", "em",
    "code", "pre", "blockquote", "table", "thead", "tbody", "tr", "th", "td", "a",
]
_ALLOWED_ATTRS = {"a": ["href", "title"]}

_env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)), trim_blocks=True, lstrip_blocks=True)

TEMPLATE_EXEC_SUMMARY = (
    "This dataset contains {n_rows} rows and {n_cols} columns. After automated cleaning and feature "
    "engineering, AutoDS-AI evaluated {n_models} models for a {problem_type} problem"
    "{target_clause}. The best-performing model was **{best_model_name}**, achieving a cross-validated "
    "score of {best_score}. {findings_clause}"
)


def build_template_executive_summary(context: dict[str, Any]) -> str:
    summary = context["summary"]
    leaderboard = context["leaderboard"]
    findings = context["eda_results"].get("notable_findings", [])
    target_clause = f" with target `{context['target_column']}`" if context.get("target_column") else ""
    findings_clause = f"Key observations: {findings[0]}" if findings else ""
    best_score = leaderboard[0]["cv_score"] if leaderboard else "N/A"
    return TEMPLATE_EXEC_SUMMARY.format(
        n_rows=summary["n_rows"],
        n_cols=summary["n_cols"],
        n_models=len(leaderboard),
        problem_type=context["problem_type"],
        target_clause=target_clause,
        best_model_name=context.get("best_model_name", "N/A"),
        best_score=best_score,
        findings_clause=findings_clause,
    )


def render_report(context: dict[str, Any], executive_summary: str | None = None) -> tuple[str, str]:
    template = _env.get_template("report.md.j2")
    exec_summary = executive_summary or build_template_executive_summary(context)
    render_context = {
        **context,
        "executive_summary": exec_summary,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "metrics_json": json.dumps(context.get("metrics", {}), indent=2, default=str),
        "shap_global": context.get("shap_results", {}).get("global_importance", []),
    }
    markdown_text = template.render(**render_context)
    raw_html = markdown_lib.markdown(markdown_text, extensions=["tables"])
    html_text = bleach.clean(raw_html, tags=_ALLOWED_TAGS, attributes=_ALLOWED_ATTRS, strip=True)
    return markdown_text, html_text


# ---------------------------------------------------------------------------
# PDF export via WeasyPrint
# ---------------------------------------------------------------------------

_PDF_WRAPPER = """\
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.5;
         max-width: 800px; margin: 40px auto; color: #111; }}
  h1 {{ font-size: 18pt; color: #1a1a2e; }}
  h2 {{ font-size: 14pt; color: #16213e; border-bottom: 1px solid #ddd; }}
  h3 {{ font-size: 12pt; color: #0f3460; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 10px; text-align: left; }}
  th {{ background: #f0f0f0; }}
  code, pre {{ background: #f8f8f8; padding: 2px 4px; border-radius: 3px;
               font-size: 9pt; }}
  pre {{ padding: 10px; overflow-x: auto; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""


def render_pdf(html_body: str) -> bytes:
    """Convert sanitized HTML body to PDF bytes via WeasyPrint."""
    try:
        from weasyprint import HTML  # lazy import
    except ImportError as exc:
        raise ImportError("weasyprint is required for PDF export: pip install weasyprint") from exc

    wrapped = _PDF_WRAPPER.format(body=html_body)
    return HTML(string=wrapped).write_pdf()


# ---------------------------------------------------------------------------
# DOCX export via python-docx
# ---------------------------------------------------------------------------

def _strip_md_inline(text: str) -> str:
    """Remove markdown inline formatting (bold, italic, backtick) for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def render_docx(markdown_text: str) -> bytes:
    """Convert markdown report text to DOCX bytes via python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise ImportError("python-docx is required for DOCX export: pip install python-docx") from exc

    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code_block = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_table = False

    def _flush_code():
        if code_lines:
            para = doc.add_paragraph("\n".join(code_lines))
            para.style = doc.styles["Normal"]
            para.runs[0].font.name = "Courier New"
            para.runs[0].font.size = Pt(9)
            code_lines.clear()

    def _flush_table():
        if len(table_rows) < 2:
            table_rows.clear()
            return
        cols = len(table_rows[0])
        tbl = doc.add_table(rows=len(table_rows) - 1, cols=cols)
        tbl.style = "Table Grid"
        # Header row (row 0 in markdown is the header, row 1 is separator)
        header_cells = table_rows[0]
        data_rows = table_rows[2:]  # skip separator row
        # Add header as a new first row
        hdr_row = tbl.rows[0]
        for i, cell_text in enumerate(header_cells[:cols]):
            cell = hdr_row.cells[i]
            cell.text = _strip_md_inline(cell_text)
            cell.paragraphs[0].runs[0].bold = True
        # Add data rows
        for r_idx, row_data in enumerate(data_rows):
            if r_idx + 1 < len(tbl.rows):
                row = tbl.rows[r_idx + 1]
            else:
                row = tbl.add_row()
            for c_idx, cell_text in enumerate(row_data[:cols]):
                row.cells[c_idx].text = _strip_md_inline(cell_text)
        table_rows.clear()

    for line in markdown_text.splitlines():
        # Code block fence
        if line.startswith("```"):
            if in_code_block:
                _flush_code()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Markdown table row
        if line.startswith("|"):
            if in_table and table_rows:
                pass  # keep accumulating
            else:
                in_table = True
                if table_rows:
                    _flush_table()
            cells = [c.strip() for c in line.strip("|").split("|")]
            table_rows.append(cells)
            continue
        else:
            if in_table:
                _flush_table()
                in_table = False

        # Headings
        if line.startswith("### "):
            doc.add_heading(_strip_md_inline(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(_strip_md_inline(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(_strip_md_inline(line[2:]), level=1)
        # List items
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(_strip_md_inline(line[2:]), style="List Bullet")
        # Horizontal rule → empty paragraph
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph()
        # Italic metadata line (e.g. _Generated ..._)
        elif line.startswith("_") and line.endswith("_"):
            para = doc.add_paragraph()
            run = para.add_run(line.strip("_"))
            run.italic = True
        # Normal paragraph (skip blank lines)
        elif line.strip():
            doc.add_paragraph(_strip_md_inline(line))

    # Flush any remaining code or table
    _flush_code()
    if in_table:
        _flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
